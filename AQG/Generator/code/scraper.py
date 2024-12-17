import os
import re
from PyPDF2 import PdfWriter, PdfReader
from tkinter import filedialog, messagebox, simpledialog, Tk, Button, Label
from pdf2docx import Converter
import pdfplumber
import fitz
from docx import Document
from docx2pdf import convert

# Fungsi untuk konversi Word ke PDF menggunakan docx2pdf
def docx_to_pdf(docx_path):
    pdf_output = os.path.splitext(docx_path)[0] + ".pdf"
    convert(docx_path, pdf_output)
    return pdf_output

# Function to delete unnecessary pages from PDF with double detection pass
def delete_pages(pdf_path, romawi_pages_end, keywords=["Daftar Pustaka", "Glosarium", "DAFTAR PUSTAKA", "GLOSARIUM"]):
    with pdfplumber.open(pdf_path) as pdf:
        writer = PdfWriter()
        total_pages = len(pdf.pages)
        detected_pages = []

        # Iterasi pertama untuk mendeteksi halaman dengan kata kunci dari belakang
        for keyword in keywords:
            for i in range(total_pages - 1, romawi_pages_end - 1, -1):
                try:
                    text = pdf.pages[i].extract_text()
                    if text and keyword.lower() in text.lower():
                        detected_pages.append(i)
                        break  # Keluar dari loop jika ditemukan
                except Exception as e:
                    print(f"Error processing page {i}: {e}")
                    continue  # Lanjutkan ke halaman berikutnya meskipun ada error

        # Jika ada halaman yang terdeteksi
        if detected_pages:
            delete_start_page = min(detected_pages)
            reader = PdfReader(pdf_path)
            for i in range(romawi_pages_end, delete_start_page):
                try:
                    writer.add_page(reader.pages[i])
                except Exception as e:
                    print(f"Error adding page {i}: {e}")
                    continue  # Lanjutkan ke halaman berikutnya

        else: 
            reader = PdfReader(pdf_path)
            for i in range(romawi_pages_end, total_pages):
                try:
                    writer.add_page(reader.pages[i])
                except Exception as e:
                    print(f"Error adding page {i}: {e}")
                    continue  # Lanjutkan ke halaman berikutnya

        output_path = f"{pdf_path.split('.')[0]}_cleaned.pdf"
        with open(output_path, "wb") as out_pdf:
            writer.write(out_pdf)

        return output_path
    
# Fungsi untuk mendeteksi TOC dan nomor halaman
def extract_chapter_titles_and_pages(pdf_path):
    chapters = []
    seen_chapters = set()
    pattern = re.compile(r'^(BAB|Tema)\s+[IVXLCDM0-9]+\.?\s+.*?(?:\.\.\.|—|\s+)+(\d+)$', re.IGNORECASE)

    with pdfplumber.open(pdf_path) as pdf:
        all_texts = []
        for page_index, page in enumerate(pdf.pages):
            try:
                text = page.extract_text()
                if not text:
                    print(f"Warning: No text found on page {page_index + 1}")
                    continue
                
                lines = text.split('\n')
                combined_lines = []
                buffer = ''
                
                for line in lines:
                    if re.match(r'^\s*[A-Z]\.\s+|^\s*\d+\.\s+', line):
                        if buffer:
                            combined_lines.append(buffer)
                            buffer = ''
                    if re.match(r'^(BAB|Tema)\s+[IVXLCDM0-9]+\.?', line, re.IGNORECASE):
                        if buffer:
                            combined_lines.append(buffer)
                        buffer = line
                    else:
                        buffer += ' ' + line
                
                if buffer:
                    combined_lines.append(buffer)

                all_texts.extend(combined_lines)
            except Exception as e:
                print(f"Error extracting text from page {page_index + 1}: {e}")
                continue

        for line in all_texts:
            clean_line = ' '.join(line.split()).strip()
            match = pattern.match(clean_line)
            if match:
                title = match.group(0).strip()
                try:
                    page_number = int(match.group(2))
                    main_identifier = re.match(r'^(BAB|Tema)\s+[IVXLCDM0-9]+', title, re.IGNORECASE).group(0)
                    if main_identifier not in seen_chapters:
                        seen_chapters.add(main_identifier)
                        chapters.append((title, page_number))
                        if len(chapters) >= 4:  # Stop after 4 chapters
                            break
                except ValueError:
                    print(f"Error converting page number for line: {clean_line}")

        if not chapters:
            print("No chapters found using the pattern.")
    return chapters

def is_chapter_heading(text):
    chapter_heading_patterns = [
        re.compile(r'^\s*(BAB|Tema|TEMA)\s+[IVXLCDM0-9]+\b', re.IGNORECASE),
        re.compile(r'^\s*(BAB|Tema|TEMA)\s+\d+\b', re.IGNORECASE)
    ]
    for pattern in chapter_heading_patterns:
        if pattern.search(text):
            return True
    return False

# Function to find chapters in a PDF with updated patterns
def detect_chapters(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        chapter_pages = []
        last_chapter_page = None
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and is_chapter_heading(text):
                if last_chapter_page is None or i - last_chapter_page > 5:
                    chapter_pages.append((i + 1, text.strip()))  # `i + 1` sudah dalam bentuk `int`
                    last_chapter_page = i

        return chapter_pages

# Function to select a chapter from a list of chapters
def select_chapter(chapters):
    chapter_string = "\n".join(
        "{}. Page {}: {}".format(i + 1, page_num, heading.splitlines()[0])
        for i, (page_num, heading) in enumerate(chapters)
    )
    
    selected_index = simpledialog.askinteger(
        "Select Chapter",
        f"Choose a chapter:\n{chapter_string}\nEnter the number of the chapter to keep:"
    )
    
    if selected_index is not None and 1 <= selected_index <= len(chapters):
        return selected_index
    else:
        messagebox.showerror("Error", "Invalid selection")
        return None

# Fungsi untuk menentukan rentang halaman setiap bab berdasarkan TOC
def determine_chapter_ranges(chapters, total_pages):
    chapter_ranges = []
    for i, (title, start_page) in enumerate(chapters):
        # Pastikan start_page adalah int
        start_page = int(start_page)
        if i + 1 < len(chapters):
            end_page = int(chapters[i + 1][1]) - 1
        else:
            end_page = total_pages
        chapter_ranges.append((title, start_page, end_page))
    
    print("\nChapter Ranges:")
    for title, start, end in chapter_ranges:
        print(f"Title: {title}, Start Page: {start}, End Page: {end}")
    
    return chapter_ranges

# Perbaiki `save_selected_chapter` agar menggunakan `chapter_ranges`
# def save_selected_chapter(pdf_path, chapter_ranges, selected_index):
#     reader = PdfReader(pdf_path)
#     writer = PdfWriter()

#     # Pastikan start_page dan end_page diambil dengan benar
#     start_page = chapter_ranges[selected_index][1] - 1  # Konversi ke indeks Python
#     end_page = chapter_ranges[selected_index][2]

#     for i in range(start_page, end_page):
#         writer.add_page(reader.pages[i])

#     output_pdf_path = os.path.join('Generator/assets/upload', f"Chapter_{selected_index}_selected.pdf")
#     with open(output_pdf_path, 'wb') as output_pdf_file:
#         writer.write(output_pdf_file)

#     return output_pdf_path
def save_selected_chapter(pdf_path, chapter_ranges, selected_index):
    reader = PdfReader(pdf_path)
    writer = PdfWriter()

    start_page = chapter_ranges[selected_index][1] - 1  # Convert to Python index
    end_page = chapter_ranges[selected_index][2]

    for i in range(start_page, end_page):
        try:
            page = reader.pages[i]
            writer.add_page(page)
        except Exception as e:
            print(f"Warning: Skipping page {i + 1} due to an error: {e}")
            continue  # Skip problematic pages and continue with the next

    output_pdf_path = os.path.join('Generator/assets/upload', f"Chapter_{selected_index}_selected.pdf")
    try:
        with open(output_pdf_path, 'wb') as output_pdf_file:
            writer.write(output_pdf_file)
    except Exception as e:
        print(f"Error writing PDF file: {e}")
        return None  # Return None or handle accordingly

    return output_pdf_path

# Function to convert a PDF to a Word document
def pdf_to_word(pdf_file, output_docx):
    try:
        # Try the standard conversion first
        cv = Converter(pdf_file)
        cv.convert(output_docx, start=0, end=None)
        cv.close()
        print("Standard PDF to Word conversion completed successfully.")
    except Exception as e:
        print(f"Warning: Error during PDF to Word conversion, switching to fallback method: {e}")
        
        # Fallback method using pdfplumber to extract text and save to Word document
        try:
            # Initialize a new Word Document
            doc = Document()
            
            # Open the PDF file with pdfplumber
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    # Extract text from each page
                    text = page.extract_text()
                    if text:  # Check if there is any text on the page
                        # Split text by newlines and add each line to the Word document
                        for line in text.split('\n'):
                            doc.add_paragraph(line)
                        doc.add_paragraph()  # Add a blank line after each page
            
            # Save the document
            doc.save(output_docx)
            print("Fallback PDF to Word conversion completed successfully.")
            
            # Call the function to remove images, tables, and footers
            remove_images_tables_footers(output_docx)
            print("Images, tables, and footers removed.")
            
        except Exception as fallback_error:
            print(f"Error during fallback PDF to Word conversion: {fallback_error}")

# Function to remove images, tables, and footers from a Word document
def remove_images_tables_footers(docx_file):
    doc = Document(docx_file)
    footer_patterns = [r"\b\d+\b\s*Kelas IX SMP/MTs Edisi Revisi\b", r"Ilmu Pengetahuan Sosial\s*\b\d+\b", 
                        r"\b\d+\b\s*Kelas VIII SMP/MTS", r"\b\d+\b\s*ILMU PENGETAHUAN SOSIAL UNTUK SMP KELAS VIII",
                        r"\b\d+\b\s*ILMU PENGETAHUAN SOSIAL UNTUK SMP/MTs KELAS VII (EDISI REVISI)", 
                        r"TEMA 01: KONDISI GEOGRAFIS DAN PELESTARIAN SUMBER DAYA ALAM\s*\b\d+\b", 
                        r"Tema 02: Kemajemukan Masyarakat Indonesia\s*\b\d+\b",
                        r"Tema 03: Nasionalisme dan Jati Diri Bangsa\s*\b\d+\b",
                        r"Tema 04: Pembangunan  Perekonomian Indonesia\s*\b\d+\b",
                        r"Tema I: Kehidupan Sosial dan Kondisi Lingkungan Sekitar\s*\b\d+\b",
                        r"Tema II: Keberagaman Lingkungan Sekitar\s*\b\d+\b",
                        r"Tema III: Potensi Ekonomi Lingkungan\s*\b\d+\b",
                        r"Tema IV: Pemberdayaan Masyarakat\s*\b\d+\b",
                        r"\b\d+\b\s*Kelas VII SMP/MTs",
                        r"\b\d+\b\s*Kelas\s*VIII\s*SMP/MTs\b",
                        r"\b\d+\b\s*Kelas\s*IX\s*SMP/MTs\b",
                        r"Ilmu\s*Pengetahuan\s*Sosial\s*\b\d+\b",
                        r"\b\d+\b\s*Kelas\s*VII\s*SMP/MTs\b",
                        r"\b\d+\b\s*ILMU\s*PENGETAHUAN\s*SOSIAL\s*UNTUK\s*SMP\s*KELAS\s*VIII\b",
                        r"\b\d+\b\s*ILMU\s*PENGETAHUAN\s*SOSIAL\s*UNTUK\s*SMP/MTs\s*KELAS\s*VII\s*\(EDISI\s*REVISI\)",
                        r"TEMA\s*01:\s*KONDISI\s*GEOGRAFIS\s*DAN\s*PELESTARIAN\s*SUMBER\s*DAYA\s*ALAM\s*\b\d+\b",
                        r"Tema\s*02:\s*Kemajemukan\s*Masyarakat\s*Indonesia\s*\b\d+\b",
                        r"Tema\s*03:\s*Nasionalisme\s*dan\s*Jati\s*Diri\s*Bangsa\s*\b\d+\b",
                        r"Tema\s*04:\s*Pembangunan\s*Perekonomian\s*Indonesia\s*\b\d+\b",
                        r"Tema\s*I:\s*Kehidupan\s*Sosial\s*dan\s*Kondisi\s*Lingkungan\s*Sekitar\s*\b\d+\b",
                        r"Tema\s*II:\s*Keberagaman\s*Lingkungan\s*Sekitar\s*\b\d+\b",
                        r"Tema\s*III:\s*Potensi\s*Ekonomi\s*Lingkungan\s*\b\d+\b",
                        r"Tema\s*IV:\s*Pemberdayaan\s*Masyarakat\s*\b\d+\b"
                        ]
    combined_footer_pattern = r'|'.join(footer_patterns)

    for para in doc.paragraphs:
        try:
            # Hapus elemen gambar dan gambar SVG
            for run in para.runs:
                if run.element.xpath(".//w:drawing") or run.element.xpath(".//w:object"):
                    run.clear()
            
            # Hapus elemen footer yang cocok
            if re.search(combined_footer_pattern, para.text):
                para.clear()
        
        except Exception as e:
            print(f"Warning: Error processing paragraph or element: {e}")
            continue  # Ignore errors and proceed

    for table in doc.tables:
        try:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = ''
        except Exception as e:
            print(f"Warning: Error clearing table: {e}")
            continue  # Ignore errors and proceed

    doc.save(docx_file)

# Function to extract text from a Word document
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = []

    for para in doc.paragraphs:
        text.append(para.text)

    return "\n".join(text)

# Function to clean extra line breaks from text
def clean_extra_linebreaks(text):
    return re.sub(r'\n\s*\n+', '\n', text)

# Function to detect and separate lists with appropriate line breaks
def detect_and_separate_list(text):
    lines = text.split("\n")
    detected_lines = []
    inside_list = False
    buffer = []
    
    list_pattern = r'^\s*(\d+\.\s+|\d+\s+|[•\-\*]\s+|[A-Z]\.\s+|[a-z]\.\s+|[ivxlcdm]+\.\s+)'
    
    for i, line in enumerate(lines):
        line = line.strip()
        if re.match(list_pattern, line):
            if buffer:
                detected_lines.extend(buffer)
                buffer = []
            detected_lines.append(line)
            inside_list = True
        else:
            if inside_list:
                buffer.append(line)
                lookahead_lines = lines[i+1:i+9]
                if not any(re.match(list_pattern, l.strip()) for l in lookahead_lines):
                    detected_lines.append("")
                    detected_lines.extend(buffer)
                    buffer = []
                    inside_list = False
            else:
                detected_lines.append(line)

    final_text = "\n".join(detected_lines)
    final_text = re.sub(r'(\d+\.\s[^\.!?]+\s*[\.!?])\s(?!\d+\.)', r'\1\n\n', final_text)
    return final_text

# Function to detect and remove task sections
def detect_task_sections_with_regex(text):
    task_keywords = [r"Aktivitas Kelompok", r"Aktivitas Individu", r"Tugas Mandiri", r"Uji Kompetensi", r"Pilihan Ganda", r"A. Pilihan Ganda", r"Esai", r"B. Esai", r"Proyek", r"Rangkuman", r"Renungkan", r"Peta Konsep", r"Uji Pemahaman Materi", r"Refleksi", r"Releksi dan Tindak Lanjut", r"releksi"]
    
    task_pattern = r"(" + "|".join(task_keywords) + r")\s*(?:(?:\n[\d\)\.]+\s+.+?)|(?:\n[A-Z]\.\s+.+?)|(?:\n[a-z]\.\s+.+?)|(?:\n\s{0,4}.+?))+?(?=\n\n|\Z)"
    cleaned_text = re.sub(task_pattern, '', text, flags=re.DOTALL | re.IGNORECASE)
    cleaned_text = re.sub(r'\n\s*\n+', '\n', cleaned_text)
    return cleaned_text

# Pastikan fungsi ini hanya melakukan pembersihan teks awal
import re

def process_manual_text_input(text):
    roman_patterns = [
        r'^\s*[IVXLCDM]+\s*Kelas\s*(IX|VIII|VII)\s*SMP/MTs\s*(Edisi\s*Revisi)?\s*$',
        r'^\s*Ilmu\s*Pengetahuan\s*Sosial\s*[IVXLCDM]+\s*$',
        r'^(?=[IVXLCDM]+\s*$)(?=[^a-zA-Z]*[IVXLCDM]+[^a-zA-Z]*$)'
    ]
    sections_to_remove = ["Daftar Pustaka", "Glosarium", "GLOSARIUM", "DAFTAR PUSTAKA"]

    lines = text.split('\n')
    roman_pages = [i for i, line in enumerate(lines) for pattern in roman_patterns if re.match(pattern, line.strip(), re.IGNORECASE)]
    if roman_pages:
        last_roman_index = roman_pages[-1]
        lines = lines[last_roman_index + 1:]

    indices = [i for i in range(len(lines) - 1, -1, -1) if any(re.search(fr'^{section}\s*$', lines[i], re.IGNORECASE) for section in sections_to_remove)]
    if indices:
        lines = lines[:min(indices)]

    cleaned_text = '\n'.join(lines)
    return cleaned_text

# Fungsi untuk mendeteksi bab/tema dari teks
def text_process_scanning(text):
    chapter_pattern = r'^\s*(BAB|Tema|TEMA)\s+[IVXLCDM0-9]+\b'
    chapters = []

    lines = text.split('\n')
    for i, line in enumerate(lines):
        if re.match(chapter_pattern, line.strip(), re.IGNORECASE):
            cleaned_line = line.strip()
            chapters.append((i, cleaned_line))

    if chapters:
        print("Debug: Detected Chapters:")
        for line_num, title in chapters:
            print(f"Line {line_num}: {title}")
    else:
        print("Debug: No chapters detected.")
    
    return chapters