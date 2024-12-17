from rouge_score import rouge_scorer
from docx import Document
import pdfplumber
import os

def calculate_rouge(input_text, reference_text):
    scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
    scores = scorer.score(reference_text, input_text)
    word_count = len(input_text.split())
    print("Scores:", scores)  # Debug: Print scores to check values
    return scores

def load_reference_text_from_upload(chapter_number):
    pdf_path = f'Generator/assets/upload/chapter_{chapter_number}_selected.pdf'
    docx_path = f'Generator/assets/upload/chapter_{chapter_number}_selected.docx'
    no_chapter_docx_path = 'Generator/assets/upload/no_chapter_selected.docx'
    no_chapter_pdf_path = 'Generator/assets/upload/no_chapter_selected.pdf'
    
    # Cek file PDF spesifik chapter
    if os.path.exists(pdf_path):
        # Ekstraksi teks dari PDF
        text = ''
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ''  # Menangani halaman kosong dengan menambahkan string kosong
        
        if not text:
            raise ValueError(f"File {pdf_path} tidak mengandung teks yang dapat diekstraksi.")
        
        return text

    elif os.path.exists(docx_path):
        # Ekstraksi teks dari Word
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ''])
        
        if not text:
            raise ValueError(f"File {docx_path} tidak mengandung teks yang dapat diekstraksi.")
        
        return text

    # Cek file `no_chapter_selected.docx`
    elif os.path.exists(no_chapter_docx_path):
        doc = Document(no_chapter_docx_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ''])
        
        if not text:
            raise ValueError(f"File {no_chapter_docx_path} tidak mengandung teks yang dapat diekstraksi.")
        
        return text

    # Cek file `no_chapter_selected.pdf`
    elif os.path.exists(no_chapter_pdf_path):
        text = ''
        with pdfplumber.open(no_chapter_pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ''  # Menangani halaman kosong dengan string kosong
        
        if not text:
            raise ValueError(f"File {no_chapter_pdf_path} tidak mengandung teks yang dapat diekstraksi.")
        
        return text

    else:
        raise FileNotFoundError("Tidak ditemukan file PDF atau Word untuk chapter yang dipilih atau 'no_chapter_selected'.")