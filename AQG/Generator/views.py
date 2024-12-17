from django.shortcuts import render, redirect
from .code.scraper import (delete_pages, extract_chapter_titles_and_pages, detect_chapters, determine_chapter_ranges,
                           save_selected_chapter, pdf_to_word, remove_images_tables_footers, extract_text_from_docx,
                           clean_extra_linebreaks, detect_and_separate_list, detect_task_sections_with_regex, 
                           process_manual_text_input, text_process_scanning)
from .code.textrank import extractive_summary, preprocess_text  
from .code.template_based import generate_questions_and_answers, load_additional_entities, analyze_text_with_labeled_entities, enrich_entities, load_templates_from_excel, match_templates, select_questions
from .code.rouge import calculate_rouge, load_reference_text_from_upload
from django.http import JsonResponse
import pdfplumber
from docx import Document
import re
import os

# Buat direktori upload jika belum ada
upload_dir = 'Generator/assets/upload'
if not os.path.exists(upload_dir):
    os.makedirs(upload_dir)

def index(request):
    # Hapus semua file di folder upload saat halaman dimuat
    for filename in os.listdir(upload_dir):
        file_path = os.path.join(upload_dir, filename)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f'Error saat menghapus file: {file_path}, error: {e}')
    
    return render(request, 'index.html')


def summary(request):
    summary_text = request.session.get('summary_text', '')  # Ambil summary_text dari sesi
    return render(request, 'summary.html', {'summary_text': summary_text})

# def generate(request):
#     summary_text = request.session.get('summary_text', '')
#     generate_q = request.session.get('questions', '')
#     return render(request, 'generate_question.html', summary_text)

def input_text(request):
    if request.method == 'POST':
        manual_text = request.POST.get('manual_text')
        if not manual_text:
            return JsonResponse({'error': 'Manual text input is required'}, status=400)

        try:
            cleaned_text = process_manual_text_input(manual_text)
            print("Debug: Cleaned text after initial processing:", cleaned_text[:500])

            chapters = text_process_scanning(cleaned_text)
            if chapters:
                request.session['chapters'] = [{"line": i, "title": title} for i, title in chapters]
                request.session['cleaned_text'] = cleaned_text

                # Simpan cleaned_text ke dalam file Word
                output_path = os.path.join(upload_dir, "no_chapter_selected.docx")
                document = Document()
                document.add_paragraph(cleaned_text)
                document.save(output_path)

                return JsonResponse({"chapters": request.session['chapters'], "success": True})
            else:
                # Kirimkan sinyal bahwa tidak ada bab ditemukan
                request.session['cleaned_text'] = cleaned_text

                # Simpan cleaned_text ke dalam file Word
                output_path = os.path.join(upload_dir, "no_chapter_selected.docx")
                document = Document()
                document.add_paragraph(cleaned_text)
                document.save(output_path)

                return JsonResponse({'no_chapters': True, 'message': 'No chapters found. Do you want to continue?'}, status=200)

        except Exception as e:
            print(f"Error in input_text function: {e}")
            return JsonResponse({'error': f'An error occurred: {e}'}, status=500)

    return JsonResponse({'error': 'Invalid request method'}, status=400)

def upload_file(request):
    if request.method == 'POST':
        uploaded_file = request.FILES.get('file_upload')
        if uploaded_file:
            file_path = os.path.join(upload_dir, uploaded_file.name)
            with open(file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)

            romawi_pages_end = int(request.POST.get('romawi_pages', 0))
            chapters = extract_chapter_titles_and_pages(file_path)
            cleaned_pdf_path = delete_pages(file_path, romawi_pages_end)

            if not chapters:
                print("TOC not detected, falling back to detect_chapters function.")
                chapters = detect_chapters(cleaned_pdf_path)

            if not chapters:
                # Simpan file hasil pembersihan untuk diproses penuh
                return JsonResponse({'no_chapters': True, 'message': 'No chapters found. Do you want to continue with the cleaned file?'}, status=200)

            with pdfplumber.open(cleaned_pdf_path) as pdf:
                total_pages = len(pdf.pages)

            chapter_ranges = determine_chapter_ranges(chapters, total_pages)
            if chapter_ranges:
                request.session['chapter_ranges'] = chapter_ranges
                request.session['file_name'] = uploaded_file.name

                chapter_data = [
                    {
                        "title": re.match(r'((Bab|Tema)\s+[IVXLCDM0-9]+)', title, re.IGNORECASE).group(0),
                        "start_page": start,
                        "end_page": end
                    }
                    for title, start, end in chapter_ranges
                ]
                return JsonResponse({"chapters": chapter_data})
            else:
                print("No chapter ranges found.")
                return JsonResponse({"error": "Chapters not found"}, status=400)

    return render(request, 'index.html')

def process_selected_chapter(request):
    if request.method == 'POST':
        selected_index = request.POST.get('selected_index')
        continue_without_chapters = request.POST.get('continue_without_chapters', False)
        cleaned_text = request.session.get('cleaned_text', '')
        file_name = request.POST.get('file_name')
        print("Selected index (backend):", selected_index)
        print("Continue without chapters (backend):", continue_without_chapters)

        if continue_without_chapters and cleaned_text:
            preprocessed_text = preprocess_text(cleaned_text)
            if not preprocessed_text:
                return JsonResponse({'error': 'Tidak ada teks yang tersedia untuk diproses.'}, status=400)
            
            try:
                # output_path = os.path.join(upload_dir, "no_chapter_selected.docx")
                # document = Document()
                # document.add_paragraph(preprocessed_text)
                # document.save(output_path)
                summary = extractive_summary(preprocessed_text, summary_ratio=0.7)
                request.session['summary_text'] = summary
                return JsonResponse({'success': True})
            except Exception as e:
                print(f"Error dalam proses ringkasan: {e}")
                return JsonResponse({'error': f'An error occurred: {e}'}, status=500)

        # Validasi selected_index untuk mencegah kesalahan
        if selected_index is None:
            return JsonResponse({'error': 'Selected index is required unless continuing without chapters.'}, status=400)

        try:
            selected_index = int(selected_index)
        except ValueError:
            return JsonResponse({'error': 'Invalid selected index format'}, status=400)

        chapter_ranges = request.session.get('chapter_ranges')
        cleaned_text = request.session.get('cleaned_text')
        chapters = request.session.get('chapters')

        if cleaned_text and chapters:
            print("Processing manual text input...")
            if selected_index >= len(chapters):
                return JsonResponse({'error': 'Indeks yang dipilih di luar rentang'}, status=400)

            start_line = chapters[selected_index]['line']
            end_line = chapters[selected_index + 1]['line'] if selected_index + 1 < len(chapters) else len(cleaned_text.split('\n'))
            selected_chapter_text = '\n'.join(cleaned_text.split('\n')[start_line:end_line])

            output_path = os.path.join(upload_dir, f"selected_chapter_{selected_index + 1}.docx")
            document = Document()
            document.add_heading(chapters[selected_index]['title'], level=1)
            document.add_paragraph(selected_chapter_text)
            document.save(output_path)

            try:
                remove_images_tables_footers(output_path)
                extracted_text = extract_text_from_docx(output_path)
                cleaned_text = clean_extra_linebreaks(extracted_text)
                text_with_linebreaks = detect_and_separate_list(cleaned_text)
                final_cleaned_text = detect_task_sections_with_regex(text_with_linebreaks)
                preprocessed_text = preprocess_text(final_cleaned_text)

                request.session['preprocessed_text'] = preprocessed_text
                summary = extractive_summary(final_cleaned_text, summary_ratio=1.0)
                request.session['summary_text'] = summary

                return JsonResponse({'success': True})
            except Exception as e:
                print(f"Error dalam pemrosesan manual text: {e}")
                return JsonResponse({'error': f'An error occurred: {e}'}, status=500)

        elif chapter_ranges:
            print("Processing PDF/Word input...")
            if selected_index >= len(chapter_ranges):
                return JsonResponse({'error': 'Indeks yang dipilih di luar rentang'}, status=400)

            cleaned_file_path = os.path.join(upload_dir, f"{os.path.splitext(file_name)[0]}_cleaned.pdf")
            if not os.path.exists(cleaned_file_path):
                print(f"File tidak ditemukan: {cleaned_file_path}")
                return JsonResponse({'error': 'File tidak ditemukan'}, status=400)

            try:
                chapter_pdf_path = save_selected_chapter(cleaned_file_path, chapter_ranges, selected_index)
                if not chapter_pdf_path:
                    return JsonResponse({'error': 'Error processing chapter PDF'}, status=500)

                output_docx = os.path.splitext(chapter_pdf_path)[0] + ".docx"
                pdf_to_word(chapter_pdf_path, output_docx)
                remove_images_tables_footers(output_docx)
                extracted_text = extract_text_from_docx(output_docx)
                cleaned_text = clean_extra_linebreaks(extracted_text)
                text_with_linebreaks = detect_and_separate_list(cleaned_text)
                final_cleaned_text = detect_task_sections_with_regex(text_with_linebreaks)
                preprocessed_text = preprocess_text(final_cleaned_text)

                request.session['preprocessed_text'] = preprocessed_text
                summary = extractive_summary(final_cleaned_text, summary_ratio=1.0)
                request.session['summary_text'] = summary

                return JsonResponse({'success': True})
            except Exception as e:
                print(f"Error dalam process_selected_chapter: {e}")
                return JsonResponse({'error': f'An error occurred: {e}'}, status=500)

        return JsonResponse({'error': 'Data tidak ditemukan untuk pemrosesan'}, status=400)

    return JsonResponse({'error': 'Metode request tidak valid'}, status=400)

def generate_questions(request):
    if request.method == 'POST':
        question_type = request.POST.get('question_type')
        num_questions = request.POST.get('num_questions', '0')  # Default ke '0' jika kosong
        
        try:
            num_questions = int(num_questions)
        except ValueError:
            return JsonResponse({'error': 'Invalid number of questions'}, status=400)

        # Cek apakah num_questions valid
        if num_questions <= 0:
            return JsonResponse({'error': 'Number of questions must be greater than 0'}, status=400)

        # Ambil teks dari sesi yang sudah diproses
        full_text = request.session.get('preprocessed_text', '')

        # Muat entitas dan template
        labeled_entities = load_additional_entities()
        entities = analyze_text_with_labeled_entities(full_text, labeled_entities)
        enriched_entities = enrich_entities(entities, labeled_entities)

        templates = load_templates_from_excel(question_type)
        matched_questions = match_templates(enriched_entities, templates)
        selected_questions = select_questions(matched_questions, num_questions * 2)

        # Generate soal dan jawaban
        questions = generate_questions_and_answers(full_text, selected_questions, num_questions, question_type, labeled_entities)

        return render(request, 'generate_question.html', {
            'summary_text': request.session.get('summary_text', ''),
            'questions': questions
        })

    return redirect('summary')

def calculate_accuracy(request):
    if request.method == 'POST':
        comparison_type = request.POST.get('comparison_type')
        system_summary = request.session.get('summary_text', '')

        if comparison_type == 'book':
            chapter_number = request.session.get('selected_chapter_number', 1)
            try:
                reference_text = load_reference_text_from_upload(chapter_number)
            except FileNotFoundError as e:
                return JsonResponse({'error': str(e)}, status=404)
        elif comparison_type == 'human':
            reference_text = request.POST.get('human_summary', '')
            if not reference_text:
                return JsonResponse({'error': 'Ringkasan manusia tidak boleh kosong'}, status=400)
        else:
            return JsonResponse({'error': 'Tipe perbandingan tidak valid'}, status=400)

        scores = calculate_rouge(system_summary, reference_text)
        print("Scores before returning:", scores)  # Debug: Print scores to check values

        response_data = {
            'scores': {
                'rouge1': {
                    'precision': scores['rouge1'].precision,
                    'recall': scores['rouge1'].recall,
                    'fmeasure': scores['rouge1'].fmeasure,
                },
                'rouge2': {
                    'precision': scores['rouge2'].precision,
                    'recall': scores['rouge2'].recall,
                    'fmeasure': scores['rouge2'].fmeasure,
                },
                'rougeL': {
                    'precision': scores['rougeL'].precision,
                    'recall': scores['rougeL'].recall,
                    'fmeasure': scores['rougeL'].fmeasure,
                }
            },
            'word_count': len(system_summary.split())
        }

        return JsonResponse(response_data)

    return JsonResponse({'error': 'Metode request tidak valid'}, status=400)